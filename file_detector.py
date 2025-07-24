# -*- coding: utf-8 -*-
import pandas as pd
import os
# pytesseract импорт с поддержкой типов через заглушки в stubs/pytesseract/__init__.pyi
import pytesseract
import cv2
import numpy as np
import pdfplumber
import PyPDF2
from docx import Document

class FileTypeDetector:
    """
    Расширенный детектор типов файлов согласно ТЗ v5.0:
    - Табличные: CSV, XLSX, XLS, TXT
    - Документы: PDF, DOC, DOCX  
    - Изображения: JPG, JPEG, PNG, BMP, TIFF (с OCR)
    - 6 основных типов источников данных
    """
    
    def __init__(self):
        # Поддерживаемые расширения файлов
        self.supported_extensions = {
            'tabular': ['.csv', '.xlsx', '.xls', '.txt'],
            'documents': ['.pdf', '.doc', '.docx'],
            'images': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
        }
        
        # Шаблоны для определения типов данных
        self.templates = {
            'happy_workers': {
                'filename_patterns': ['hw', 'happy', 'workers', 'автомат'],
                'columns': ['order number', 'machine code', 'creation time', 
                           'paying time', 'brewing time', 'delivery time',
                           'order price', 'goods name', 'payment status', 'refund time'],
                'threshold': 0.7  # Повышен согласно ТЗ
            },
            'vendhub': {
                'filename_patterns': ['report', 'vendhub', 'внутренний', 'учет'],
                'columns': ['order number', 'goods id', 'goods name', 'order price',
                           'machine code', 'machine category', 'payment type',
                           'order time', 'order resource', 'username'],
                'threshold': 0.7  # Повышен согласно ТЗ
            },
            'fiscal_bills': {
                'filename_patterns': ['fiscal', 'bills', 'чек', 'фискал'],
                'columns': ['fiscal_check_number', 'fiscal_time', 'amount', 
                           'taxpayer_id', 'cash_register_id', 'shift_number'],
                'threshold': 0.8  # Согласно ТЗ
            },
            'payme': {
                'filename_patterns': ['payme', 'пейми'],
                'columns': ['transaction_id', 'transaction_time', 'amount',
                           'masked_pan', 'merchant_id', 'terminal_id', 'status'],
                'threshold': 0.8  # Согласно ТЗ
            },
            'click': {
                'filename_patterns': ['click', 'клик'],
                'columns': ['transaction_id', 'transaction_time', 'amount',
                           'card_number', 'merchant_id', 'service_id', 'status'],
                'threshold': 0.8  # Согласно ТЗ
            },
            'uzum': {
                'filename_patterns': ['uzum', 'узум'],
                'columns': ['transaction_id', 'transaction_time', 'amount',
                           'masked_pan', 'shop_id', 'merchant_id', 'status'],
                'threshold': 0.8  # Согласно ТЗ
            }
        }
        
        # Настройка OCR
        self.ocr_languages = ['rus', 'eng']
    
    def detect_type(self, file_path):
        """
        Расширенное определение типа файла согласно ТЗ v5.0:
        1. Определение категории файла (табличный, документ, изображение)
        2. Обработка в зависимости от категории
        3. Определение типа данных по содержимому
        """
        filename = os.path.basename(file_path).lower()
        file_ext = os.path.splitext(filename)[1].lower()
        
        # Определяем категорию файла
        file_category = self._get_file_category(file_ext)
        
        if file_category == 'unsupported':
            print(f"Unsupported file format: {file_ext}")
            return 'unknown'
        
        # Шаг 1: Определение по имени файла
        filename_match = self._detect_by_filename(filename)
        if filename_match:
            print(f"File type detected by filename: {filename_match}")
            return filename_match
        
        # Шаг 2: Обработка в зависимости от категории
        try:
            if file_category == 'tabular':
                return self._process_tabular_file(file_path, file_ext)
            elif file_category == 'documents':
                return self._process_document_file(file_path, file_ext)
            elif file_category == 'images':
                return self._process_image_file(file_path)
            else:
                return 'unknown'
                
        except Exception as e:
            print(f"Error detecting file type: {e}")
            return 'unknown'
    
    def _get_file_category(self, file_ext):
        """Определение категории файла по расширению"""
        for category, extensions in self.supported_extensions.items():
            if file_ext in extensions:
                return category
        return 'unsupported'
    
    def _process_tabular_file(self, file_path, file_ext):
        """Обработка табличных файлов"""
        df = None
        
        if file_ext in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path, nrows=5)
        elif file_ext == '.csv':
            # Пробуем разные кодировки для CSV
            for encoding in ['utf-8', 'cp1251', 'windows-1251', 'iso-8859-1']:
                try:
                    df = pd.read_csv(file_path, nrows=5, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
        elif file_ext == '.txt':
            # Пробуем разные разделители для TXT
            for sep in ['\t', ',', ';', '|']:
                try:
                    df = pd.read_csv(file_path, nrows=5, sep=sep, encoding='utf-8')
                    if len(df.columns) > 1:  # Если нашли разделитель
                        break
                except (UnicodeDecodeError, pd.errors.EmptyDataError, pd.errors.ParserError):
                    continue
        
        if df is None:
            print(f"Could not read tabular file: {file_path}")
            return 'unknown'
        
        return self._classify_by_columns(df.columns)
    
    def _process_document_file(self, file_path, file_ext):
        """Обработка документов (PDF, DOC, DOCX)"""
        try:
            text_content = ""
            
            if file_ext == '.pdf':
                text_content = self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                text_content = self._extract_docx_text(file_path)
            elif file_ext == '.doc':
                # Для .doc файлов нужны дополнительные библиотеки
                print("DOC files not fully supported yet")
                return 'unknown'
            
            # Пытаемся найти табличные данные в тексте
            if text_content:
                return self._classify_by_text_content(text_content)
            
            return 'unknown'
            
        except Exception as e:
            print(f"Error processing document: {e}")
            return 'unknown'
    
    def _process_image_file(self, file_path):
        """Обработка изображений с OCR"""
        try:
            # Предобработка изображения
            img = cv2.imread(file_path)
            if img is None:
                return 'unknown'
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Улучшение качества для OCR
            gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
            gray = cv2.GaussianBlur(gray, (5, 5), 0)
            gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # OCR
            text = pytesseract.image_to_string(gray, lang='+'.join(self.ocr_languages))
            
            if text.strip():
                return self._classify_by_text_content(text)
            
            return 'unknown'
            
        except Exception as e:
            print(f"Error processing image with OCR: {e}")
            return 'unknown'
    
    def _extract_pdf_text(self, file_path):
        """Извлечение текста из PDF"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Сначала пытаемся извлечь таблицы
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table[:5]:  # Первые 5 строк
                                if row:
                                    text += " ".join([str(cell) for cell in row if cell]) + "\n"
                    else:
                        # Если таблиц нет, извлекаем обычный текст
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text[:1000]  # Первые 1000 символов
                    
                    if len(text) > 500:  # Достаточно для анализа
                        break
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
        
        return text
    
    def _extract_docx_text(self, file_path):
        """Извлечение текста из DOCX"""
        text = ""
        try:
            doc = Document(file_path)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows[:5]:  # Первые 5 строк
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " ".join(row_text) + "\n"
            
            # Если таблиц нет, извлекаем из параграфов
            if not text.strip():
                for paragraph in doc.paragraphs[:10]:  # Первые 10 параграфов
                    text += paragraph.text + "\n"
                    
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
        
        return text
    
    def _classify_by_columns(self, columns):
        """Классификация по названиям колонок"""
        file_columns = [str(col).lower().strip() for col in columns]
        
        best_match = None
        best_score = 0
        
        for file_type, config in self.templates.items():
            template_columns = [col.lower() for col in config['columns']]
            
            # Подсчет совпадений по колонкам
            matches = 0
            for template_col in template_columns:
                for file_col in file_columns:
                    if template_col in file_col or any(word in file_col for word in template_col.split()):
                        matches += 1
                        break
            
            score = matches / len(template_columns) if template_columns else 0
            
            if score >= config['threshold'] and score > best_score:
                best_match = file_type
                best_score = score
        
        if best_match:
            print(f"File type detected by columns: {best_match} (score: {best_score:.2f})")
        else:
            print(f"No matching template found. Available columns: {file_columns}")
        
        return best_match or 'unknown'
    
    def _classify_by_text_content(self, text):
        """Классификация по текстовому содержимому"""
        text_lower = text.lower()
        
        # Ищем ключевые слова для каждого типа
        keywords = {
            'happy_workers': ['order number', 'machine code', 'creation time', 'paying time'],
            'vendhub': ['goods id', 'machine category', 'order resource', 'username'],
            'fiscal_bills': ['fiscal', 'чек', 'taxpayer', 'налогоплательщик'],
            'payme': ['payme', 'пейми', 'transaction id', 'masked_pan'],
            'click': ['click', 'клик', 'service id'],
            'uzum': ['uzum', 'узум', 'shop id']
        }
        
        best_match = None
        best_score = 0
        
        for file_type, words in keywords.items():
            matches = sum(1 for word in words if word in text_lower)
            score = matches / len(words)
            
            if score > best_score and score >= 0.3:  # Минимум 30% совпадений
                best_match = file_type
                best_score = score
        
        if best_match:
            print(f"File type detected by text content: {best_match} (score: {best_score:.2f})")
        
        return best_match or 'unknown'
    
    def _detect_by_filename(self, filename: str) -> str:
        """Определение типа файла по имени"""
        filename = filename.lower()
        
        for file_type, config in self.templates.items():
            for pattern in config['filename_patterns']:
                if pattern.lower() in filename:
                    return file_type
        
        return 'unknown'
